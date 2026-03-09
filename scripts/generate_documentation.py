"""
Генерира Word документация за платформата за анализ на техника при гребане.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text):
    for block in text.split('\n\n'):
        if block.strip():
            p = doc.add_paragraph(block.strip())
            p.paragraph_format.space_after = Pt(6)

def create_documentation():
    doc = Document()
    
    # Заглавие
    title = doc.add_heading('Уеб базирана платформа за анализ на техника при гребане', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.add_run('Документация на дипломния проект').italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1. Въведение
    add_heading(doc, '1. Въведение', 1)
    add_para(doc, """Платформата представлява уеб-базирана система за анализ на техника при гребане чрез компютърно зрение и обработка на видео данни. Разработена е с практическа насоченост за използване от треньори и състезатели при обективна оценка на техниката и проследяване на прогреса.""")
    
    add_para(doc, """Основните цели на проекта са:
• Позволяване на качване и обработка на тренировъчни видеа
• Извличане на ключови точки от движението чрез MediaPipe Pose
• Анализ на биомеханични параметри (ротация на торса, drive/recovery, честота, симетрия)
• Сравняване с референтни модели на правилна техника
• Визуализация на метриките и предоставяне на препоръки за подобрение""")

    # 2. Функционалности
    add_heading(doc, '2. Функционалности на системата', 1)
    
    add_heading(doc, '2.1. Анализ на видео', 2)
    add_para(doc, """Системата приема тренировъчни видеа в формати MP4, AVI, MOV, MKV и WebM. Използва се MediaPipe Pose Landmarker за детекция на скелета – 33 ключови точки на тялото. За анализ на гребане се използват 11 точки: нос, рамене, лакти, китки, ханш и колена.""")
    
    add_heading(doc, '2.2. Биомеханични метрики', 2)
    add_para(doc, """Системата изчислява следните метрики:

Ротация на торса – ъгълът между линията рамене-ханш и вертикалата. Отражава амплитудата на движение между catch и finish. Оптимални стойности: 25–50°.

Drive/Recovery съотношение – съотношението между фазата на задвижване и фазата на възстановяване. При правилна техника е 1:2 до 1:2.5.

Честота на загребвания (SPM) – броят загребвания в минута. За тренировки е типично 18–36 SPM.

Симетрия – оценка на баланса между лявата и дясната страна чрез сравнение на разстоянията рамо–лакът.""")
    
    add_heading(doc, '2.3. Референтен модел', 2)
    add_para(doc, """Платформата позволява обучение от набор референтни видеа с правилна техника. Потребителят качва няколко видеа, след което системата изгражда статистически модел (средно, стандартно отклонение) за всяка метрика. При анализа на нови видеа резултатите се сравняват с този модел вместо с фиксирани стойности.""")
    
    add_heading(doc, '2.4. Чатбот-съветник', 2)
    add_para(doc, """Интегриран е чатбот за съвети по техника на гребане, спортна подготовка и изготвяне на тренировъчни планове. Поддържа се OpenAI API и Abacus.ai RouteLLM. При наличие на резултати от анализ контекстът се подава автоматично за персонализирани препоръки.""")

    # 3. Технологичен стек
    add_heading(doc, '3. Технологичен стек', 1)
    add_para(doc, """• Backend: Python, FastAPI, Uvicorn
• Компютърно зрение: OpenCV, MediaPipe Tasks
• Анализ: NumPy, SciPy
• Чатбот: OpenAI API / Abacus.ai RouteLLM
• Frontend: HTML, CSS, JavaScript""")

    # 4. Инсталация и стартиране
    add_heading(doc, '4. Инсталация и стартиране', 1)
    
    add_heading(doc, '4.1. Изисквания', 2)
    add_para(doc, """• Python 3.8 или по-нова версия
• Интернет връзка (за MediaPipe модел и опционално чатбот)""")
    
    add_heading(doc, '4.2. Стъпки за инсталация', 2)
    add_para(doc, """1. Отворете терминал в папката на проекта
2. Преминете в backend: cd backend
3. Създайте виртуална среда: python -m venv venv
4. Активирайте я: .\\venv\\Scripts\\Activate.ps1 (Windows)
5. Инсталирайте зависимости: pip install -r requirements.txt
6. Стартирайте сървъра: python run.py""")
    
    add_heading(doc, '4.3. Достъп', 2)
    add_para(doc, """Отворете уеб браузър на адрес: http://localhost:8000""")

    # 5. Конфигурация
    add_heading(doc, '5. Конфигурация', 1)
    
    add_heading(doc, '5.1. Чатбот (опционално)', 2)
    add_para(doc, """За AI чатбот създайте файл backend/.env:

За OpenAI:
OPENAI_API_KEY=sk-вашият-ключ

За Abacus.ai:
ABACUS_API_KEY=вашият-ключ
ABACUS_BASE_URL=https://routellm.abacus.ai/v1
ABACUS_MODEL=route-llm

Без ключ чатботът дава основни предварително зададени съвети.""")
    
    add_heading(doc, '5.2. Препоръки за видео', 2)
    add_para(doc, """• Ъгъл: страничен изглед е най-подходящ за ротация на торса
• Резолюция: 720p или 1080p
• Честота: 24–30 fps
• Дължина: 15–60 секунди
• Гребецът да е изцяло в кадъра""")

    # 6. Структура на проекта
    add_heading(doc, '6. Структура на проекта', 1)
    add_para(doc, """RowingAnalysis/
├── backend/              Python API и анализ
│   ├── app/              FastAPI приложение, роутери
│   ├── analysis/         Pose extraction, биомеханика, reference model
│   ├── chat/             Системен промпт за чатбота
│   └── config.py         Конфигурация на пътища
├── frontend/             Уеб интерфейс (HTML, CSS, JS)
├── uploads/              Качени видеа
├── outputs/              Резултати от анализ
└── reference_videos/     Референтни видеа за обучение""")

    # 7. API endpoints
    add_heading(doc, '7. API endpoints', 1)
    add_para(doc, """• POST /api/upload/ – качване на видео
• POST /api/analyze/{video_id} – стартиране на анализ
• GET /api/results/{video_id} – получаване на резултати
• GET/POST /api/reference/videos – управление на референтни видеа
• POST /api/reference/train – обучение на модел
• GET /api/reference/model – статус на модела
• POST /api/chat/ – чат съобщения""")

    doc.add_paragraph()
    add_para(doc, '— Край на документацията —')
    
    return doc

if __name__ == "__main__":
    from pathlib import Path
    doc = create_documentation()
    out_dir = Path(__file__).resolve().parent.parent
    out_path = out_dir / "Documentation_Rowing_Platform.docx"
    doc.save(str(out_path))
    print("Saved:", str(out_path))
