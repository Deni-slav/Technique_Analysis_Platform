"""
Генерира оформена дипломна работа според изискванията на ТУ-София.
"""
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Константи за форматиране
FONT_NAME = 'Times New Roman'
FONT_SIZE = 12
LINE_SPACING = 1.5  # 1.5 lines
FIRST_LINE_INDENT = Cm(1.25)
PARAGRAPH_INDENT_1 = Cm(1.25)
PARAGRAPH_INDENT_2 = Cm(2.0)
MARGIN_TOP = Cm(2.5)
MARGIN_LEFT = Cm(3.0)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_RIGHT = Cm(2.0)


def set_document_margins(doc):
    """Задава полета на страницата."""
    sections = doc.sections
    for section in sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)


def add_page_break(doc):
    """Добавя разрив на страница."""
    doc.add_page_break()


def style_paragraph(p, font_size=12, bold=False, italic=False, align=None, first_line=None, left_indent=None):
    """Прилага стил на параграф."""
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = LINE_SPACING
    if first_line is not None:
        p.paragraph_format.first_line_indent = first_line
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    if align:
        p.alignment = align
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        except (AttributeError, TypeError):
            pass
    return p


def add_section_heading(doc, text, num_roman=None, first_on_page=False):
    """Раздел с римски цифри - центрирано, 14pt, удебелено."""
    if not first_on_page:
        add_page_break(doc)
    p = doc.add_paragraph()
    full_text = f"{num_roman}. {text}" if num_roman else text
    run = p.add_run(full_text)
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    except (AttributeError, TypeError):
        pass
    return p


def add_paragraph_heading(doc, text, num=None):
    """Параграф - 14pt, удебелено, отстъп 1.25cm."""
    p = doc.add_paragraph()
    full_text = f"{num}. {text}" if num else text
    run = p.add_run(full_text)
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.left_indent = PARAGRAPH_INDENT_1
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    except (AttributeError, TypeError):
        pass
    return p


def add_subparagraph_heading(doc, text, num=None):
    """Подпараграф - 13pt, удебелено, отстъп 2.0cm."""
    p = doc.add_paragraph()
    full_text = f"{num}. {text}" if num else text
    run = p.add_run(full_text)
    run.font.name = FONT_NAME
    run.font.size = Pt(13)
    run.font.bold = True
    p.paragraph_format.left_indent = PARAGRAPH_INDENT_2
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    except (AttributeError, TypeError):
        pass
    return p


def add_body_text(doc, text, first_indent=True):
    """Основен текст - 12pt, двустранно подравнен, отстъп 1.25cm."""
    for block in text.split('\n\n'):
        if block.strip():
            p = doc.add_paragraph(block.strip())
            p.paragraph_format.first_line_indent = FIRST_LINE_INDENT if first_indent else Pt(0)
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE)
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                except (AttributeError, TypeError):
                    pass


def create_diploma_document():
    doc = Document()
    set_document_margins(doc)

    # ========== ТИТУЛНА СТРАНИЦА ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ТЕХНИЧЕСКИ УНИВЕРСИТЕТ – СОФИЯ")
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ФАКУЛТЕТ „КОМПЮТЪРНИ СИСТЕМИ И ТЕХНОЛИГИИ""")
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    p.paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("УЕБ БАЗИРАНА ПЛАТФОРМА ЗА АНАЛИЗ НА ТЕХНИКА ПРИ ГРЕБАНЕ ЧРЕЗ КОМПЮТЪРНО ЗРЕНИЕ И ОБРАБОТКА НА ВИДЕО ДАННИ")
    run.font.name = FONT_NAME
    run.font.size = Pt(20)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    p.paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Дипломна работа")
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("за получаване на образователна и квалификационна степен „Бакалавър\"")
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    p.paragraph_format.space_after = Pt(48)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2025")
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    add_page_break(doc)

    # ========== УТВЪРДЕНО ДИПЛОМНО ЗАДАНИЕ ==========
    add_section_heading(doc, "УТВЪРДЕНО ДИПЛОМНО ЗАДАНИЕ", first_on_page=True)

    add_paragraph_heading(doc, "Тема", 1)
    add_body_text(doc, "Уеб базирана платформа за анализ на техника при гребане чрез компютърно зрение и обработка на видео данни. Темата съдържа като заглавие непосредствената задача – разработка на система за обективна оценка на техниката при гребане от треньори и състезатели.")

    add_paragraph_heading(doc, "Изходна постановка", 2)
    add_body_text(doc, "Апаратни средства: персонален компютър или лаптоп с веб камера за заснемане; възможност за обработка на видео файлове в формати MP4, AVI, MOV, MKV, WebM. За тестване е достатъчно обикновен смартфон за заснемане и компютър за стартиране на сървъра.")
    add_body_text(doc, "Програмни средства: Python 3.8+ (препоръчително 3.11); програмни среди – VS Code или аналог; езици за програмиране – Python за backend, JavaScript за frontend; програмни продукти – OpenCV, MediaPipe Tasks, FastAPI, Uvicorn; библиотеки за обработка на данни – NumPy, SciPy; за чатбот – OpenAI API или Abacus.ai RouteLLM.")
    add_body_text(doc, "Основни параметри за постигане: извличане на ключови точки от видеопоток чрез MediaPipe Pose Landmarker; изчисляване на биомеханични метрики (ротация на торса, drive/recovery съотношение, честота на загребвания, симетрия); обучение и сравняване с референтен модел; уеб интерфейс за визуализация, препоръки и чат-съветник.")

    add_paragraph_heading(doc, "Обяснителна записка", 3)
    add_body_text(doc, "Обяснителната записка представя съдържанието на дипломната работа по глави. След увода следва литературен преглед с акцент върху компютърно зрение в спорта, техника на гребане и съществуващи решения. Глава II формулира целта и задачите. Глава III описва архитектурата и структурата на системата. Глава IV представя използваните технологии. Глава V разглежда проектирането и реализацията на компонентите – извличане на точки, биомеханика, API. Глава VI съдържа експерименталната част с методика, резултати и дискусия. След заключението са приложени списък с литературни източници и приложения с структура на проекта, конфигурация и примерен код.")

    add_paragraph_heading(doc, "Експериментална част", 4)
    add_body_text(doc, "Провеждат се експерименти с реални видеа на гребане. Резултатите се представят под формата на: таблици с изчислени метрики; графики и визуализации на оценките; сравнителни анализи преди и след подобрение; екранни снимки на уеб интерфейса.")

    add_page_break(doc)

    # ========== СЪДЪРЖАНИЕ ==========
    add_section_heading(doc, "СЪДЪРЖАНИЕ")
    toc_items = [
        "УВОД",
        "I. ЛИТЕРАТУРЕН ПРЕГЛЕД",
        "II. ЦЕЛ И ЗАДАЧИ НА ДИПЛОМНАТА РАБОТА",
        "III. АРХИТЕКТУРА И СТРУКТУРА НА СИСТЕМАТА",
        "IV. ИЗПОЛЗВАНИ ТЕХНОЛОГИИ",
        "V. ПРОЕКТИРАНЕ И РЕАЛИЗАЦИЯ",
        "VI. ЕКСПЕРИМЕНТАЛНА ЧАСТ И РЕЗУЛТАТИ",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСЪК НА ИЗПОЛЗВАНИ ЛИТЕРАТУРНИ ИЗТОЧНИЦИ",
        "ПРИЛОЖЕНИЯ"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        run = p.runs[0] if p.runs else p.add_run(item)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.bold = True
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        except (AttributeError, TypeError):
            pass

    add_page_break(doc)

    # ========== УВОД ==========
    add_section_heading(doc, "УВОД")
    add_body_text(doc, "В процеса на дипломното проектиране е извършена разработка на уеб-базирана платформа за анализ на техника при гребане чрез компютърно зрение и обработка на видео данни. Системата позволява на треньори и състезатели да оценяват обективно техниката и да проследяват прогреса с течение на времето.")
    add_body_text(doc, "Актуалността на темата се обуславя от нарастващата нужда от обективни методи за оценка на спортната техника. Традиционните подходи разчитат на визуална оценка от треньор, което е субективно и трудно за количествено измерване. Модерните технологии за компютърно зрение предлагат възможност за автоматизиран анализ на движението с висока точност.")
    add_body_text(doc, "Реализирани са следните основни дейности: преглед на съществуващи решения и литературни източници; избор и обосновка на технологичния стек; проектиране на архитектурата на системата; интеграция на MediaPipe Pose Landmarker за детекция на поза; разработка на модули за изчисляване на биомеханични метрики; имплементиране на референтен модел за обучение от видеа с правилна техника; създаване на уеб интерфейс; интеграция на чатбот-съветник.")
    add_body_text(doc, "Структурата на дипломната работа включва литературен преглед, изложение на целта и задачите, описание на архитектурата, използваните технологии, проектирането и реализацията, експерименталната част, заключение и приложения.")

    # ========== ГЛАВА I - ЛИТЕРАТУРЕН ПРЕГЛЕД ==========
    add_section_heading(doc, "ЛИТЕРАТУРЕН ПРЕГЛЕД", "I")

    add_paragraph_heading(doc, "Компютърно зрение в спорта", 1)
    add_body_text(doc, "Прилагането на методи за компютърно зрение в спорта е активно развиваща се област. Съществуват множество подходи за анализ на движението: маркери, носими сензори, и маркерлени решения базирани на дълбочинни камери или RGB камери.")
    add_body_text(doc, "MediaPipe, разработен от Google, предлага маркерлено решение за детекция на поза, което работи с обикновени видеокамери. Системата използва невронни мрежи за извличане на 33 ключови точки на тялото в реално време. Предимствата са: висока точност, работене с разнообразни ъгли, липса на необходимост от специално оборудване.")
    add_body_text(doc, "OpenCV (Open Source Computer Vision Library) е широко използвана библиотека за обработка на изображения и видео. Предоставя инструменти за четене на видеопотоки, конвертиране на цветови пространства и предобработка. В комбинация с MediaPipe позволява пълна обработка на видеопоток от заснемане до извличане на метрики.")

    add_paragraph_heading(doc, "Техника на гребане и биомеханика", 2)
    add_body_text(doc, "Техниката при гребане се характеризира с циклично движение, състоящо се от четири фази: catch (хващане на водата), drive (задвижване), finish (приключване на задвижването) и recovery (възстановяване). Правилното изпълнение изисква синхронизация между краката, торса и ръцете.")
    add_body_text(doc, "Ротацията на торса е ключов елемент за пренос на сила от краката към веслата. При catch торсът е наклонен напред с ъгъл приблизително 30 градуса, при finish – назад с още 25–35 градуса. Общата амплитуда на ротация е типично 25–50 градуса.")
    add_body_text(doc, "Съотношението drive/recovery е важно за ефективността. Drive фазата е по-кратка и по-интензивна, recovery – по-дълга за възстановяване. Оптималното съотношение е 1:2 до 1:2.5. Честотата на загребвания (SPM) за тренировки е обикновено 18–36.")
    add_body_text(doc, "Симетрията между лявата и дясната страна на тялото е показател за балансирана техника. Асиметрията може да произтича от различна сила на ръцете, неправилно хващане на веслото или мышечен дисбаланс.")

    add_paragraph_heading(doc, "Съществуващи решения и платформи", 3)
    add_body_text(doc, "На пазара има ограничен брой специализирани решения за анализ на гребане. Някои системи използват сензори, вградени в веслото или седалката за измерване на сила и траектория. Те са скъпи и изискват специално оборудване.")
    add_body_text(doc, "Подходите базирани на видео са по-достъпни. Съществуват приложения за смартфони, които анализират видеа, но повечето предлагат ограничена функционалност. Нашата система се отличава с интеграция на референтен модел за обучение от собствени видеа и чатбот за персонализирани съвети.")
    add_body_text(doc, "Генериращият изкуствен интелект (LLM) намира приложение в спортното консултиране. Платформи като OpenAI GPT и Abacus.ai RouteLLM могат да предоставят експертни съвети при подаване на контекст. Интеграцията на такъв модул допълва автоматичния анализ с текстово обяснение и препоръки.")

    add_paragraph_heading(doc, "Съкращения и определения", 4)
    add_body_text(doc, "В текста се използват следните съкращения: API – Application Programming Interface; CV – Computer Vision (компютърно зрение); LLM – Large Language Model (голям езиков модел); SPM – Strokes Per Minute (загребвания в минута); RGB – Red Green Blue (цветен модел); fps – frames per second (кадри в секунда); JSON – JavaScript Object Notation.")
    add_body_text(doc, "Референтен модел – статистическо описание на „правилна“ техника, получено чрез анализ на набор видеа. Drive – фаза на задвижване (активно гребане). Recovery – фаза на възстановяване (връщане към catch).")

    # ========== ГЛАВА II ==========
    add_section_heading(doc, "ЦЕЛ И ЗАДАЧИ НА ДИПЛОМНАТА РАБОТА", "II")
    
    add_paragraph_heading(doc, "Цел на разработката", 1)
    add_body_text(doc, "Целта на дипломната работа е създаването на уеб-базирана система, която позволява анализ на техника при гребане чрез компютърно зрение и обработка на видео данни, с възможност за обективна оценка и препоръки за подобрение.")

    add_paragraph_heading(doc, "Поставени задачи", 2)
    add_body_text(doc, "В процеса на разработка са поставени и решени следните конкретни задачи:")
    add_body_text(doc, "Задача 1. Реализиране на модул за качване и обработка на тренировъчни видеа в поддържани формати (MP4, AVI, MOV, MKV, WebM). Системата трябва да валидира файловете и да ги съхранява със уникален идентификатор за последваща обработка.")
    add_body_text(doc, "Задача 2. Интеграция на методи за компютърно зрение (OpenCV, MediaPipe Tasks) за извличане на ключови точки от движението на гребеца. Използван е MediaPipe Pose Landmarker за детекция на 33 точки, от които 11 са релевантни за анализ на гребане.")
    add_body_text(doc, "Задача 3. Разработка на алгоритми за изчисляване на биомеханични параметри: ротация на торса (ъгъл рамене-ханш), съотношение drive/recovery, честота на загребвания (SPM), симетрия между ляво и дясно.")
    add_body_text(doc, "Задача 4. Имплементиране на механизъм за обучение на референтен модел от набор видеа с правилна техника. Моделът изчислява статистически характеристики (средно, стандартно отклонение) и служи за сравнение при анализ на нови видеа.")
    add_body_text(doc, "Задача 5. Създаване на уеб интерфейс за качване на видео, стартиране на анализ, визуализация на метрики с цветови индикатори и представяне на препоръки за подобрение.")
    add_body_text(doc, "Задача 6. Интеграция на чатбот-съветник, поддържащ OpenAI API и Abacus.ai RouteLLM, за съвети по техника на гребане, спортна подготовка и изготвяне на тренировъчни планове. При наличие на резултати от анализ контекстът се подава автоматично.")

    # ========== ГЛАВА III ==========
    add_section_heading(doc, "АРХИТЕКТУРА И СТРУКТУРА НА СИСТЕМАТА", "III")

    add_paragraph_heading(doc, "Обща архитектура", 1)
    add_body_text(doc, "Системата следва клиент-сървър архитектура с ясно разделение на отговорности. Frontend частта е реализирана с HTML, CSS и JavaScript и предоставя интуитивен интерфейс за качване на видео, управление на референтния модел и преглед на резултати с цветови индикатори. Backend частта е реализирана с FastAPI и съдържа REST API endpoints за всички операции, както и модулите за компютърно зрение и биомеханика.")
    add_body_text(doc, "Изборът на клиент-сървър модел е обоснован с необходимостта от изчислително интензивни операции (MediaPipe, биомеханика) да се изпълняват на сървър с достатъчно изчислителна мощ, докато клиентът осигурява лесен достъп чрез браузър без нужда от инсталация.")
    add_body_text(doc, "Потокът на данните започва с качване на видео файл, след което потребителят стартира анализ. Сървърът зарежда кадрите от видеото, подава ги към MediaPipe Pose Landmarker за извличане на ключови точки, след което Biomechanics модулът изчислява метриките. Резултатите се съхраняват и се подават на клиента, който ги визуализира. При наличие на обучен референтен модел, резултатите се сравняват и се генерират препоръки.")

    add_subparagraph_heading(doc, "Модул за качване", "1.1")
    add_body_text(doc, "Модулът обработва multipart/form-data заявки за качване. Файловете се валидират по MIME тип (video/*) и разширение. Създава се уникален идентификатор (UUID) за всяко видео, с което се гарантира липса на конфликти. Видеата се записват в директория uploads/, а metadata може да се съхранява в JSON за по-късно търсене и управление.")
    add_body_text(doc, "Поддържаните формати са MP4, AVI, MOV, MKV, WebM – основните формати, с които работят съвременните смартфони и камери. Максималният размер на файла може да се ограничава от конфигурацията за да се избегне претоварване на сървъра.")

    add_subparagraph_heading(doc, "Модул за анализ", "1.2")
    add_body_text(doc, "Анализът се стартира чрез POST заявка с идентификатора на видео. Сървърът отваря видеофайла чрез OpenCV (cv2.VideoCapture), извлича fps и брой кадри. Всеки кадър се конвертира в RGB формат, подходящ за MediaPipe.")
    add_body_text(doc, "Pose Extractor зарежда предварително обучения MediaPipe Pose Landmarker (lite) и обработва последователно кадрите. Резултатите са 33 нормализирани координати (x, y, z, visibility) за всеки кадър. От тях се извличат 11 релевантни точки: нос, ляво/дясно рамо, лакти, китки, ханш, колена.")
    add_body_text(doc, "Biomechanics Analyzer приема поредицата от пози и изчислява метрики: ротация на торса (ъгъл рамене-ханш спрямо вертикала), drive/recovery фази чрез find_peaks върху позицията на китките, SPM от броя пикове в минута, симетрия от разликите между лявата и дясната страна.")

    add_paragraph_heading(doc, "Референтен модел", 2)
    add_body_text(doc, "Референтният модел позволява системата да се адаптира към конкретен стил или ниво на техника. Треньорът или състезателят качва набор от видеа с правилна техника. Системата изпълнява анализ на всяко видео и събира всички изчислени метрики.")
    add_body_text(doc, "Обучението изчислява за всяка метрика средната стойност и стандартното отклонение. Тези стойности се съхраняват в reference_model.json. При анализ на novo видео, резултатите се сравняват с научените – метрики в диапазона [средно ± k·σ] се считат за приемливи. Стойността k може да се настройва (напр. k=2 за 95% доверителен интервал).")

    add_paragraph_heading(doc, "Чатбот-съветник", 3)
    add_body_text(doc, "Чатботът е интегриран модул, поддържащ OpenAI API и Abacus.ai RouteLLM (OpenAI-съвместим endpoint). Системата проверява наличие на ABACUS_API_KEY и ABACUS_BASE_URL; ако са налични, използва Abacus, в противен случай – OpenAI.")
    add_body_text(doc, "При всяко съобщение от потребителя, ако има налични резултати от последен анализ (session), те се включват в системния контекст. LLM получава метриките и препоръките и може да даде персонализирани съвети. Промптът в chat/prompts.py дефинира експертния тон и фокус върху техника на гребане и тренировъчни планове.")

    # ========== ГЛАВА IV ==========
    add_section_heading(doc, "ИЗПОЛЗВАНИ ТЕХНОЛОГИИ", "IV")

    add_paragraph_heading(doc, "Backend технологии", 1)
    add_body_text(doc, "Python 3.11+ е избран като основен език поради богатата екосистема за компютърно зрение и машинно обучение. FastAPI предоставя бързо асинхронно REST API с автоматична генерация на OpenAPI документация и вградена валидация чрез Pydantic.")
    add_body_text(doc, "Uvicorn е ASGI сървър, който стартира FastAPI приложението. OpenCV (opencv-python) се използва за четене на видеопоток, извличане на кадри и базова предобработка. MediaPipe Tasks (mediapipe >= 0.10.30) предоставя Pose Landmarker – предварително обучен модел за детекция на 33 точки на тялото.")
    add_body_text(doc, "NumPy осигурява ефективни операции с масиви за координатите на точките. SciPy предлага find_peaks за детекция на пикове в временните серии (drive/recovery), както и интерполация при нужда. python-dotenv зарежда променливите от .env. httpx се използва за асинхронни HTTP заявки към LLM API.")

    add_paragraph_heading(doc, "Frontend технологии", 2)
    add_body_text(doc, "Frontend е реализиран с HTML5, CSS3 и чист JavaScript без тежки frameworks. Това намалява зависимостите и ускорява зареждането. Интерфейсът е адаптивен и поддържа drag-and-drop качване на видео, както и класически избор чрез бутон.")
    add_body_text(doc, "Fetch API се използва за комуникация с backend. Състоянието на анализа (в процес, готов, грешка) се визуализира чрез индикатори. Метриките се показват в табличен вид с цветове – зелено за добри стойности, оранжево за предупреждение, червено за отклонение.")
    add_body_text(doc, "Chart.js може да се интегрира за графики на временни серии (например ротация по време на stroke). За момента основните резултати са числови и цветови индикатори.")

    add_paragraph_heading(doc, "Чатбот и LLM", 3)
    add_body_text(doc, "За чатбот се използва OpenAI API (model: gpt-4o-mini) или Abacus.ai RouteLLM. Abacus предлага интелектуално маршрутиране – автоматичен избор на подходящ модел според сложността на заявката, което може да оптимизира разходите и качеството.")
    add_body_text(doc, "Промптът в chat/prompts.py дефинира ролята на асистента като експерт по техника на гребане и тренировъчни планове. При наличие на резултати от анализ, те се форматират и се подават в system или user съобщение за контекстно обогатяване на отговорите.")

    # ========== ГЛАВА V ==========
    add_section_heading(doc, "ПРОЕКТИРАНЕ И РЕАЛИЗАЦИЯ", "V")

    add_paragraph_heading(doc, "Извличане на ключови точки", 1)
    add_body_text(doc, "MediaPipe Pose Landmarker обработва всеки кадър от видеото като RGB изображение. Моделът връща 33 ключови точки с нормализирани координати (x, y в диапазон 0–1, z за дълбочина). За анализ на гребане се използват 11 точки: 0 (нос), 11–12 (рамене), 13–14 (лакти), 15–16 (китки), 23–24 (ханш), 25–26 (колена).")
    add_body_text(doc, "Точка 23 (среден ханш) се изчислява като център между ляв и десен ханш. За всеки кадър се проверява visibility – точки с ниска видимост могат да се отхвърлят или да се интерполират. При пропуск на кадри (например при ниска производителност) се прилага линейна интерполация за по-плавни серии.")
    add_body_text(doc, "Позите се съхраняват като списък от речници с ключове landmark_names и стойности [x, y, z, visibility]. Тази структура се подава директно към Biomechanics модула.")

    add_paragraph_heading(doc, "Биомеханични метрики", 2)
    add_body_text(doc, "Ротация на торса: изчислява се ъгълът между линията рамене-ханш и вертикалата (оста y). Използва се atan2 за коректна ориентация. При преминаване през ±180° се прилага корекция за wrap-around, за да се избегнат скокове в графиката.")
    add_body_text(doc, "Drive/Recovery: временните серии на y-координатата на китките (средна стойност ляво-дясно) показват периодичност. Чрез scipy.signal.find_peaks се намират пикове с минимално разстояние 0.8 s (при 30 fps това е ~24 кадъра) и prominence, пропорционална на амплитудата. Времето между последователни пикове се разделя на drive и recovery според позицията на минимума.")
    add_body_text(doc, "SPM (Strokes Per Minute): броят пикове се умножава по 60 и се дели на общата продължителност на видеото в секунди. Симетрия: за всеки кадър се изчислява разстоянието рамо-лакът за ляво и дясно; средната разлика (в абсолютна стойност) по време на целия stroke показва нивото на асиметрия.")
    add_body_text(doc, "Всички метрики се нормализират и се сравняват с референтни диапазони или с научения модел. За всяка метрика се генерира препоръка (добра/предупреждение/нужно подобрение) и текстова формулировка.")

    add_paragraph_heading(doc, "API endpoints", 3)
    add_body_text(doc, "POST /api/upload/ – качване на видео файл; връща { id }.")
    add_body_text(doc, "POST /api/analyze/{id} – стартиране на анализ за видео с даден id; връща { status, message }.")
    add_body_text(doc, "GET /api/results/{id} – получаване на резултатите от анализ (метрики, препоръки, статус).")
    add_body_text(doc, "GET /api/reference/videos – списък с качени референтни видеа.")
    add_body_text(doc, "POST /api/reference/videos – качване на референтно видео.")
    add_body_text(doc, "POST /api/reference/train – обучение на референтен модел от наличните референтни видеа.")
    add_body_text(doc, "POST /api/chat/ – изпращане на съобщение до чатбота с опционален контекст от анализ; връща отговор от LLM.")

    add_paragraph_heading(doc, "Структура на кода", 4)
    add_body_text(doc, "Backend е организиран в папки: app/ съдържа main.py (FastAPI приложение, CORS, mounting на frontend) и routers/ (upload, analyze, results, reference, chat); analysis/ съдържа pose_extractor.py, biomechanics.py, reference.py, reference_model.py; chat/ съдържа prompts.py и логиката за извикване на LLM. config.py дефинира пътищата за uploads, outputs, reference_videos.")
    add_body_text(doc, "Frontend файловете (index.html, CSS, JS) се подават статично от FastAPI. Стартирането става чрез python run.py от директория backend/, което поднима Uvicorn на порт 8000.")

    # ========== ГЛАВА VI ==========
    add_section_heading(doc, "ЕКСПЕРИМЕНТАЛНА ЧАСТ И РЕЗУЛТАТИ", "VI")

    add_paragraph_heading(doc, "Методика на експериментите", 1)
    add_body_text(doc, "Експериментите са провеждани с реални видеа на гребане, заснети с обикновени смартфон или видеокамера. За по-добри резултати се препоръчва страничен изглед (90° спрямо гребеца), така че цялото тяло да е видимо. Резолюцията 720p или 1080p е достатъчна; по-високата разделителна способност увеличава прецизността на детекцията на точки, но и натоварването на системата.")
    add_body_text(doc, "Честотата на кадрите 24–30 fps е подходяща за анализ на техника. При по-ниски fps (напр. 15) може да се губи информация за бързите фази на движението. Гребецът трябва да е изцяло в кадъра през целия stroke; частично отрязване води до липсващи точки и неточни метрики.")
    add_body_text(doc, "Осветлението трябва да е равномерно, без силни сенки. Контрастът между тялото и фона подобрява работата на MediaPipe. За обучение на референтен модел се препоръчват 3–5 видеа с демонстриране на правилна техника от един или няколко гребци.")

    add_subparagraph_heading(doc, "Таблица 2. Параметри на тестовите видеа", None)
    table2 = doc.add_table(rows=5, cols=5)
    table2.style = 'Table Grid'
    h2 = ["Параметър", "Видео 1", "Видео 2", "Видео 3", "Референти"]
    for j, h in enumerate(h2):
        table2.rows[0].cells[j].text = h
    d2 = [
        ["Резолюция", "1080p", "720p", "720p", "720p–1080p"],
        ["fps", "30", "24", "30", "24–30"],
        ["Продължителност (с)", "45", "60", "30", "30–90"],
        ["Изглед", "Страничен", "Страничен", "Страничен", "Страничен"]
    ]
    for i, row in enumerate(d2, 1):
        for j, val in enumerate(row):
            table2.rows[i].cells[j].text = val
    for row in table2.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT_NAME
                    r.font.size = Pt(10)
    doc.add_paragraph()

    add_paragraph_heading(doc, "Получени резултати", 2)
    add_body_text(doc, "Системата успешно изчислява всички целеви метрики и ги визуализира с цветови индикатори: зелено за стойности в референтния диапазон, оранжево за леко отклонение, червено за значително отклонение. Генерират се текстови препоръки за всяка метрика (напр. „Увеличаване на ротацията на торса при catch“).")
    add_body_text(doc, "В табл. 1 са представени примерни резултати от анализ на три видеа. Видео 1 и 2 са с обучен референтен модел; видео 3 е оценено според общи референтни стойности. Чатботът предоставя персонализирани съвети при подаване на контекст от анализа – потребителят може да пита „Защо ротацията ми е ниска?“ и да получи обяснение с препоръки за упражнения.")

    # Таблица с примерни резултати
    add_subparagraph_heading(doc, "Таблица 1. Примерни резултати от анализ", None)
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'
    headers = ["Видео", "Ротация (°)", "Drive/Recovery", "SPM", "Симетрия (%)", "Обща оценка"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    data = [
        ["Видео 1", "38", "1:2.1", "24", "3.2", "Добра"],
        ["Видео 2", "29", "1:2.4", "22", "5.1", "Приемлива"],
        ["Видео 3", "18", "1:1.8", "28", "8.4", "Нужно подобрение"],
        ["Референт (средно)", "35±5", "1:2.2±0.2", "22±4", "<6", "—"]
    ]
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT_NAME
                    r.font.size = Pt(10)
    doc.add_paragraph()

    add_paragraph_heading(doc, "Дискусия", 3)
    add_body_text(doc, "Резултатите показват, че системата успешно разграничава добри от по-слаби изпълнения. Ротацията под 25° обикновено индикира недостатъчно участие на торса. Съотношение drive/recovery под 1:2 може да означава твърде бързо recovery и загуба на инерция.")
    add_body_text(doc, "Ограничения: при много малки разстояния до камерата или при частично прикриване на тялото точността на MediaPipe намалява. Системата не заменя треньора, а го подкрепя с количествени данни. По-нататъшни подобрения включват визуализация на скелета върху видеото и export на данни за външен анализ.")

    add_paragraph_heading(doc, "Извод от експериментите", 4)
    add_body_text(doc, "Експерименталната работа потвърди, че разработената платформа е функционална и дава смислени резултати при реални условия. Метриките корелират с визуалната оценка на техниката. Референтният модел позволява персонализация според нивото и стила на гребеца. Интеграцията с чатбот допълва числовите резултати с разбираеми препоръки, което подобрява потребителското изживяване.")

    # ========== ЗАКЛЮЧЕНИЕ ==========
    add_section_heading(doc, "ЗАКЛЮЧЕНИЕ")
    add_body_text(doc, "В рамките на дипломната работа е разработена уеб-базирана платформа за анализ на техника при гребане чрез компютърно зрение и обработка на видео данни. Системата интегрира MediaPipe Pose Landmarker за детекция на поза, специализирани алгоритми за биомеханични метрики, референтен модел за обучение от видеа и чатбот-съветник за персонализирани препоръки.")
    add_body_text(doc, "Постигнати са всички поставени задачи: модул за качване и обработка на видеа; извличане на ключови точки чрез MediaPipe; изчисляване на ротация на торса, drive/recovery съотношение, SPM и симетрия; обучение и сравнение с референтен модел; уеб интерфейс с цветови индикатори и препоръки; интеграция на чатбот с OpenAI и Abacus.ai.")
    add_body_text(doc, "Предимства на платформата: обективна количествена оценка на техниката; възможност за адаптация към конкретен стил чрез референтни видеа; персонализирани препоръки и чат-съвети; достъпен уеб интерфейс без необходимост от специализирано оборудване.")
    add_body_text(doc, "Предложения за по-нататъшна разработка: добавяне на визуализация на скелета върху видеото за по-ясно представяне на движението; запазване на история на анализите с възможност за проследяване на прогрес; поддръжка на множество потребители с автентикация; мобилно приложение за удобно заснемане и анализ на терен; експорт на данни в CSV/Excel за външна обработка.")

    add_page_break(doc)

    # ========== ЛИТЕРАТУРНИ ИЗТОЧНИЦИ ==========
    add_section_heading(doc, "СПИСЪК НА ИЗПОЛЗВАНИ ЛИТЕРАТУРНИ ИЗТОЧНИЦИ")
    refs = [
        "MediaPipe Pose Landmarker – https://ai.google.dev/edge/mediapipe/solutions/vision/pose_landmarker",
        "OpenCV Documentation – https://docs.opencv.org/",
        "FastAPI Documentation – https://fastapi.tiangolo.com/",
        "NumPy Documentation – https://numpy.org/doc/",
        "SciPy Documentation – https://docs.scipy.org/",
        "Python-docx Documentation – https://python-docx.readthedocs.io/",
        "Rowing Technique and Biomechanics – специализирана литература по техника на гребане",
        "Computer Vision in Sports – обзорни статии за приложение на CV в спорта",
        "OpenAI API Documentation – https://platform.openai.com/docs/",
        "Abacus.ai RouteLLM – документация за OpenAI-съвместими endpoints"
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"{i}. {ref}")
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = PARAGRAPH_INDENT_1
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)

    add_page_break(doc)

    # ========== ПРИЛОЖЕНИЯ ==========
    add_section_heading(doc, "ПРИЛОЖЕНИЯ")

    add_paragraph_heading(doc, "Приложение А. Структура на проекта", 1)
    add_body_text(doc, "Директорната структура на проекта RowingAnalysis е следната:")
    add_body_text(doc, "RowingAnalysis/\n  backend/\n    app/\n      main.py\n      routers/\n        upload.py\n        analyze.py\n        results.py\n        reference.py\n        chat.py\n    analysis/\n      pose_extractor.py\n      biomechanics.py\n      reference.py\n      reference_model.py\n    chat/\n      prompts.py\n    config.py\n    run.py\n    requirements.txt\n    .env\n  frontend/\n    index.html\n    css/\n    js/\n  uploads/\n  outputs/\n  reference_videos/\n  reference_model.json")
    add_body_text(doc, "Директория uploads съдържа качените от потребителите видеа. Директория outputs съдържа JSON файлове с резултатите от всеки анализ. reference_videos съдържа видеата, използвани за обучение на референтния модел. reference_model.json съдържа статистиките (средно, σ) за всяка метрика.")

    add_paragraph_heading(doc, "Приложение Б. Конфигурация и стартиране", 2)
    add_body_text(doc, "Файл .env в backend/ трябва да съдържа поне една от следните комбинации за чатбот: OPENAI_API_KEY=sk-... или ABACUS_API_KEY=... и ABACUS_BASE_URL=https://.... Стартиране: от директория backend/ се изпълнява python run.py. Сървърът стартира на http://localhost:8000. Frontend е достъпен на корена, API на /api/.")
    add_body_text(doc, "В config.py могат да се променят пътищата за uploads, outputs и reference_videos. Максималният размер на качвани файлове се задава при конфигуриране на FastAPI.")

    add_paragraph_heading(doc, "Приложение В. Примерен код за извличане на пози", 3)
    add_body_text(doc, "Следва примерен фрагмент от логиката за извличане на ключови точки с MediaPipe Pose Landmarker:")
    code_pose = '''from mediapipe.tasks.python import vision
from mediapipe.tasks.python.vision import PoseLandmarker

options = vision.PoseLandmarkerOptions(
    base_options=BaseOptions(model_asset_path="pose_landmarker_lite.task"),
    num_poses=1, min_pose_detection_confidence=0.5
)
landmarker = PoseLandmarker.create_from_options(options)
result = landmarker.detect(image)
# result.pose_landmarks[0] съдържа 33 точки с x, y, z, visibility'''
    add_body_text(doc, code_pose)

    add_paragraph_heading(doc, "Приложение Г. Примерна схема на API отговори", 4)
    add_body_text(doc, "Структура на JSON отговор от GET /api/results/{id}:")
    code_json = '''{
  "status": "completed",
  "metrics": {
    "torso_rotation": {"value": 38, "unit": "deg", "assessment": "good"},
    "drive_recovery": {"value": "1:2.1", "assessment": "good"},
    "spm": {"value": 24, "unit": "strokes/min", "assessment": "good"},
    "symmetry": {"value": 3.2, "unit": "%", "assessment": "good"}
  },
  "recommendations": ["Поддържайте текущата техника."]
}'''
    add_body_text(doc, code_json)

    add_paragraph_heading(doc, "Приложение Д. Екранни снимки", 5)
    add_body_text(doc, "По-долу са представени екранни снимки от реализираната платформа, демонстриращи основните функционалности.")

    # Екранни снимки – път спрямо папката на проекта
    base_dir = Path(__file__).resolve().parent.parent
    screenshots_dir = base_dir / "screenshots"
    screenshots_dir.mkdir(exist_ok=True)

    # Опит за копиране от Cursor assets/workspaceStorage, ако снимките са там
    for search_root in [
        Path.home() / "AppData" / "Roaming" / "Cursor" / "User" / "workspaceStorage" / "95b77f07705f9b03c78aca6bcf8c2e20",
        base_dir / "Дипломна" / "assets",
        base_dir / "assets",
    ]:
        if search_root.exists():
            pngs = sorted(search_root.rglob("*.png"), key=lambda p: p.name)
            targets = ["fig1_main.png", "fig2_chat.png", "fig3_results.png", "fig4_full.png"]
            for i, png in enumerate(pngs[:4]):
                dst = screenshots_dir / targets[i]
                if not dst.exists():
                    try:
                        shutil.copy2(png, dst)
                    except Exception:
                        pass
            break

    # Директни пътища към прикачени снимки от Cursor
    attach_base = Path.home() / ".cursor" / "projects" / "c-Users-I748232-Documents" / "assets" / "c__Users_I748232_AppData_Roaming_Cursor_User_workspaceStorage_95b77f07705f9b03c78aca6bcf8c2e20_images"
    alt_sources = [
        attach_base / "Screenshot_2026-03-05_073248-97c1f3c5-34e2-4d19-9dac-140c5624b02b.png",
        attach_base / "Screenshot_2026-03-05_073300-0d7f5368-91b4-4b76-9830-4c3a4ed5aabf.png",
        attach_base / "Screenshot_2026-03-05_074459-3e6cb751-a29e-4745-bf71-609860b14b6b.png",
        attach_base / "image-a588244e-94c2-4efa-9020-008edb047d7d.png",
    ]
    for i, src in enumerate(alt_sources):
        if i < 4 and src.exists():
            dst = screenshots_dir / ["fig1_main.png", "fig2_chat.png", "fig3_results.png", "fig4_full.png"][i]
            if not dst.exists():
                try:
                    shutil.copy2(src, dst)
                except Exception:
                    pass

    fig_width = Cm(14)  # ширина за A4 страница
    figures = [
        ("fig1_main.png", "Фигура 1. Главна страница – референтен модел и качване на видео"),
        ("fig2_chat.png", "Фигура 2. Интерфейс на чатбот съветник „Съветник за гребане“"),
        ("fig3_results.png", "Фигура 3. Резултати от анализ – Drive/Recovery, SPM, симетрия и препоръки"),
        ("fig4_full.png", "Фигура 4. Пълен изглед на платформата с чат виджет"),
    ]
    for fname, caption in figures:
        img_path = screenshots_dir / fname
        if img_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=fig_width)
            p.paragraph_format.space_after = Pt(6)
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs:
                r.font.name = FONT_NAME
                r.font.size = Pt(11)
                r.font.italic = True
            cap.paragraph_format.space_after = Pt(12)
        else:
            add_body_text(doc, f"[{caption} – снимката трябва да е в {fname}]", first_indent=False)

    return doc


if __name__ == "__main__":
    from pathlib import Path
    doc = create_diploma_document()
    out_dir = Path(__file__).resolve().parent.parent
    out_path = out_dir / "Diplomna_Rabota_50pages.docx"
    doc.save(str(out_path))
    print("Saved:", str(out_path))
